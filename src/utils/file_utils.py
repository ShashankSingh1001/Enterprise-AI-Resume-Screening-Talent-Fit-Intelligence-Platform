import os
import sys
from pathlib import Path
from typing import Dict, Optional
from datetime import datetime

import PyPDF2
import pdfplumber
from docx import Document

from src.logging import get_logger
from src.exceptions import DataValidationError

logger = get_logger(__name__)

# Allowed file formats
ALLOWED_FORMATS = ["pdf", "docx", "doc", "txt"]
DEFAULT_MAX_SIZE_MB = 10


def read_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file.

    Raises:
        DataValidationError: If PDF reading fails
    """
    try:
        logger.info(f"Reading PDF file: {file_path}")

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        text_content = []

        # Preferred: pdfplumber
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)

            logger.info(f"Extracted {len(text_content)} pages using pdfplumber")
            return "\n".join(text_content)

        except Exception as plumber_error:
            logger.warning(f"pdfplumber failed, fallback to PyPDF2: {plumber_error}")

            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)

            logger.info(f"Extracted {len(text_content)} pages using PyPDF2")
            return "\n".join(text_content)

    except Exception as e:
        logger.error(f"Failed to read PDF: {file_path}")
        raise DataValidationError(
            message=f"Failed to read PDF: {file_path}",
            error_detail=sys.exc_info()
        )


def read_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file.

    Raises:
        DataValidationError
    """
    try:
        logger.info(f"Reading DOCX file: {file_path}")

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        doc = Document(file_path)
        text_content = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)

        logger.info(f"Extracted {len(text_content)} sections from DOCX")
        return "\n".join(text_content)

    except Exception as e:
        logger.error(f"Failed to read DOCX: {file_path}")
        raise DataValidationError(
            message=f"Failed to read DOCX: {file_path}",
            error_detail=sys.exc_info()
        )


def read_txt(file_path: str, encoding: str = "utf-8") -> str:
    """
    Read plain text file.

    Raises:
        DataValidationError
    """
    try:
        logger.info(f"Reading text file: {file_path}")

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            with open(file_path, "r", encoding=encoding) as file:
                return file.read()

        except UnicodeDecodeError:
            logger.warning(f"{encoding} failed, using latin-1")
            with open(file_path, "r", encoding="latin-1") as file:
                return file.read()

    except Exception as e:
        logger.error(f"Failed to read text file: {file_path}")
        raise DataValidationError(
            message=f"Failed to read text file: {file_path}",
            error_detail=sys.exc_info()
        )


def validate_file_format(filename: str, allowed_formats: Optional[list] = None) -> bool:
    """
    Validate file extension.
    """
    allowed_formats = allowed_formats or ALLOWED_FORMATS
    extension = filename.split(".")[-1].lower()

    is_valid = extension in allowed_formats
    if not is_valid:
        logger.warning(f"Invalid file format: {extension} (allowed: {allowed_formats})")

    return is_valid


def validate_file_size(file_path: str, max_size_mb: int = DEFAULT_MAX_SIZE_MB) -> bool:
    """
    Validate file size.

    Raises:
        DataValidationError
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        size_mb = os.path.getsize(file_path) / (1024 * 1024)
        return size_mb <= max_size_mb

    except Exception as e:
        logger.error(f"Failed to validate file size: {file_path}")
        raise DataValidationError(
            message=f"Failed to validate file size: {file_path}",
            error_detail=sys.exc_info()
        )


def save_uploaded_file(
    file_content: bytes,
    filename: str,
    upload_dir: str = "data/uploads"
) -> str:
    """
    Save uploaded file safely.

    Raises:
        DataValidationError
    """
    try:
        upload_path = Path(upload_dir)
        upload_path.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        name, ext = os.path.splitext(filename)
        unique_name = f"{name}_{timestamp}{ext}"

        file_path = upload_path / unique_name
        with open(file_path, "wb") as f:
            f.write(file_content)

        logger.info(f"File saved: {file_path}")
        return str(file_path)

    except Exception as e:
        logger.error(f"Failed to save uploaded file: {filename}")
        raise DataValidationError(
            message=f"Failed to save uploaded file: {filename}",
            error_detail=sys.exc_info()
        )


def get_file_metadata(file_path: str) -> Dict[str, object]:
    """
    Extract file metadata.

    Raises:
        DataValidationError
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        stat = os.stat(file_path)

        return {
            "filename": os.path.basename(file_path),
            "extension": os.path.splitext(file_path)[1],
            "size_bytes": stat.st_size,
            "size_mb": stat.st_size / (1024 * 1024),
            "created_at": datetime.fromtimestamp(stat.st_ctime),
            "modified_at": datetime.fromtimestamp(stat.st_mtime),
            "absolute_path": os.path.abspath(file_path),
        }

    except Exception as e:
        logger.error(f"Failed to extract metadata: {file_path}")
        raise DataValidationError(
            message=f"Failed to extract metadata: {file_path}",
            error_detail=sys.exc_info()
        )


def delete_file(file_path: str) -> bool:
    """
    Delete file safely.

    Raises:
        DataValidationError
    """
    try:
        if not os.path.exists(file_path):
            logger.warning(f"File not found for deletion: {file_path}")
            return False

        os.remove(file_path)
        logger.info(f"File deleted: {file_path}")
        return True

    except Exception as e:
        logger.error(f"Failed to delete file: {file_path}")
        raise DataValidationError(
            message=f"Failed to delete file: {file_path}",
            error_detail=sys.exc_info()
        )
